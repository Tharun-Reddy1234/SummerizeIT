import os, requests, csv
from flask import Flask, render_template, request, send_file
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from docx import Document
from io import BytesIO

load_dotenv()
app = Flask(__name__)
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")

def extract_text_from_url(url):
    r = requests.get(url)
    soup = BeautifulSoup(r.text, "html.parser")
    return ' '.join([p.get_text() for p in soup.find_all("p")])

def summarize_text(text):
    payload = {
        "model": "google/gemma-3-27b-it:free",  # Or use gemma-3-12b-it / gemma-3-4b-it
        "messages": [
            {"role": "user", "content": f"Summarize this content:\n{text[:8000]}"}
        ]
    }
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {OPENROUTER_API_KEY}"
    }

    res = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers=headers,
        json=payload
    )

    try:
        response_json = res.json()
        return response_json["choices"][0]["message"]["content"]
    except Exception as e:
        return f"❌ OpenRouter response parsing failed: {str(e)}"

@app.route("/", methods=["GET", "POST"])
def home():
    if request.method == "POST":
        url = request.form["website_url"]
        raw_text = extract_text_from_url(url)
        summary = summarize_text(raw_text)
        return render_template("result.html", summary=summary)
    return render_template("index.html")

@app.route("/download", methods=["POST"])
def download():
    summary = request.form["summary"]
    format = request.form["format"]

    if format == "word":
        doc = Document()
        doc.add_heading("Website Summary", 0)
        doc.add_paragraph(summary)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name="summary.docx")

    elif format == "csv":
        buffer = BytesIO()
        writer = csv.writer(buffer)
        writer.writerow(["Summary"])
        writer.writerow([summary])
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name="summary.csv", mimetype="text/csv")

    return "Invalid format selected", 400

if __name__ == "__main__":
    app.run(debug=True)
